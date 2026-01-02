from core.action.action_framework.registry import action

@action(
    name="create word file",
    description="This action creates a Word (.docx) file at the specified file path using the provided Markdown content. It converts Markdown to a Word document using python-docx, supporting headings (levels 1-6), paragraphs, bold, italics, bullet lists, and numbered lists. If the content is empty or invalid, the action will return an appropriate error. This action is production-ready and handles file permission errors gracefully.",
    mode="CLI",
    input_schema={
        "file_path": {
            "type": "string",
            "example": "/home/user/documents/my_file.docx",
            "description": "The full path where the new Word file will be created. Make sure the directory exists and is writable by the agent."
        },
        "content": {
            "type": "string",
            "example": "# Heading 1\\n\\nThis is **bold** text.\\n\\n- Item 1\\n- Item 2\\n\\n1. First\\n2. Second",
            "description": "The Markdown-formatted content to be converted into a Word document. Supports headings (#, ##, etc.), paragraphs, bullet lists (-, *), numbered lists (1., 2.), bold (**text**), and italics (*text*)."
        }
    },
    output_schema={
        "status": {
            "type": "string",
            "example": "success",
            "description": "'success' if the file was created, 'error' otherwise."
        },
        "path": {
            "type": "string",
            "example": "/home/user/documents/my_file.docx",
            "description": "The path to the newly created Word document."
        },
        "message": {
            "type": "string",
            "example": "Permission denied.",
            "description": "Error message, present only if status is 'error'."
        }
    },
    requirement=["Document", "Pt", "BeautifulSoup", "markdown2", "docx", "bs4"],
    test_payload={
            "file_path": "/home/user/documents/my_file.docx",
            "content": "# Heading 1\\n\\nThis is **bold** text.\\n\\n- Item 1\\n- Item 2\\n\\n1. First\\n2. Second",
            "simulated_mode": True
    }
)
def create_word_file(input_data: dict) -> dict:
    import json
    import sys
    import subprocess
    import importlib
    import os
    import re

    def _ensure(pkg: str) -> None:
        try:
            importlib.import_module(pkg)
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "--quiet"])

    # Ensure required packages are installed
    [_ensure(_pkg) for _pkg in ("python-docx", "markdown2", "beautifulsoup4")]

    from docx import Document
    from docx.shared import Pt
    from bs4 import BeautifulSoup
    import markdown2

    def _add_paragraph_with_formatting(doc, element):
        text = ""
        runs = []

        for item in element.descendants:
            if isinstance(item, str):
                text += item
            elif item.name in ("strong", "b"):
                runs.append(("bold", item.get_text()))
            elif item.name in ("em", "i"):
                runs.append(("italic", item.get_text()))

        if not runs:
            doc.add_paragraph(text.strip())
        else:
            p = doc.add_paragraph()
            cursor = 0
            for fmt, run_text in runs:
                before = text.find(run_text, cursor)
                if before > cursor:
                    p.add_run(text[cursor:before])
                run = p.add_run(run_text)
                if fmt == "bold":
                    run.bold = True
                elif fmt == "italic":
                    run.italic = True
                cursor = before + len(run_text)
            if cursor < len(text):
                p.add_run(text[cursor:])

    simulated_mode = input_data.get('simulated_mode', False)
    
    file_path = str(input_data.get("file_path", "")).strip()
    content = str(input_data.get("content", "")).strip()

    if not file_path:
        return {
            "status": "error",
            "path": "",
            "message": "The 'file_path' field is required."
        }

    if not content:
        return {
            "status": "error",
            "path": "",
            "message": "The 'content' field is required."
        }
    
    if simulated_mode:
        # Return mock result for testing
        return {
            "status": "success",
            "path": file_path
        }

    try:
        html_content = markdown2.markdown(content)
        soup = BeautifulSoup(html_content, "html.parser")
        doc = Document()

        for element in soup.descendants:
            if element.name:
                tag = element.name.lower()
                if tag.startswith("h") and tag[1:].isdigit():
                    level = int(tag[1])
                    heading_text = element.get_text(strip=True)
                    doc.add_heading(heading_text, level=level)
                elif tag == "p":
                    _add_paragraph_with_formatting(doc, element)
                elif tag == "ul":
                    for li in element.find_all("li"):
                        doc.add_paragraph(li.get_text(strip=True), style="List Bullet")
                elif tag == "ol":
                    for li in element.find_all("li"):
                        doc.add_paragraph(li.get_text(strip=True), style="List Number")

        doc.save(file_path)
        return {
            "status": "success",
            "path": file_path
        }
    except Exception as e:
        return {
            "status": "error",
            "path": "",
            "message": str(e)
        }
